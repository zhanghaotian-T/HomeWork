import docx


class DocEdit(object):
    def __init__(self):
        self.document = self.document_open()

    def document_open(self):
        try:
            f = open(r'.\\Reports\\TestReport.docx', 'rb')
            document = docx.Document(f)
            f.close()
        except FileNotFoundError:
            document = docx.Document()
            document.save(r'.\\Reports\\TestReport.docx')
        return document

    def add_paragraph(self, args):
        for key in args.keys():
            new_paragraph = self.document.add_paragraph('The senary ' + str(key) + 'count is ' + str(args[key]))
        self.document.save(r'.\\Reports\\TestReport.docx')


if __name__ == '__main__':
    doc = DocEdit()
    maindocument = doc.document_open()

